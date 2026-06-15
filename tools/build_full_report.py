from __future__ import annotations

import os
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
IMG_DIR = OUT_DIR / "images"
REPORT_PATH = OUT_DIR / "序知_软件工程实验报告.docx"

BLUE = "1F4E79"
MID_BLUE = "5B9BD5"
LIGHT_BLUE = "DDEBF7"
PALE_BLUE = "EAF3F8"
GRID_BLUE = "9CC2E5"
TEXT = "1F1F1F"
MUTED = "666666"


def pc(color: str) -> str:
    """把 Word 用的十六进制颜色转换成 PIL 可识别的颜色。"""
    if color.startswith("#"):
        return color
    return f"#{color}"


def font_path() -> str:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    return ""


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def ensure_dirs() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    IMG_DIR.mkdir(exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID_BLUE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_para(doc: Document, text: str = "", style: str | None = None, align=None, bold=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_caption(doc: Document, text: str, kind: str = "图") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=10, color=TEXT)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F8FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=TEXT)
    doc.add_paragraph()


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, size=10.5, color=TEXT, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths_cm)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], MID_BLUE)
        header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, size=10, color="FFFFFF", bold=True)

    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            if row_index % 2 == 0:
                set_cell_shading(cells[idx], PALE_BLUE)
            else:
                set_cell_shading(cells[idx], "FFFFFF")
            para = cells[idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or len(value) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(value)
            set_run_font(run, size=9.5, color=TEXT)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("序知 AI 个人体系化知识库软件工程实验报告")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    r2 = footer.add_run(" 页")
    set_run_font(r2, size=9, color=MUTED)


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, font, fill=None, max_width=220, line_gap=6, center=False):
    if fill is None:
        fill = pc(TEXT)
    x, y = xy
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            test = current + ch
            if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tx = x + (max_width - (bbox[2] - bbox[0])) / 2 if center else x
        draw.text((tx, y), line, font=font, fill=fill)
        y += bbox[3] - bbox[1] + line_gap


def arrow(draw, start, end, fill="#333333", width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)]
    else:
        pts = [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
    draw.polygon(pts, fill=fill)


def save_image(name: str, image: Image.Image) -> Path:
    path = IMG_DIR / name
    image.save(path)
    return path


def draw_box(draw, box, label, fill="#FFFFFF", outline="#333333", title=False):
    draw.rounded_rectangle(box, radius=8, fill=fill, outline=outline, width=2)
    font = pil_font(24 if title else 18)
    x1, y1, x2, y2 = box
    draw_wrapped(draw, (x1 + 12, y1 + 14), label, font, max_width=x2 - x1 - 24, center=True)


def make_usecase() -> Path:
    img = Image.new("RGB", (1500, 900), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    title_font = pil_font(28)
    d.text((560, 35), "序知系统用例图", font=title_font, fill=pc(BLUE))
    d.rectangle((260, 100, 1220, 810), outline=pc(BLUE), width=3)
    d.text((285, 118), "序知 AI 个人体系化知识库", font=pil_font(22), fill=pc(BLUE))

    actors = {"访客": (110, 260), "登录用户": (110, 520), "AI 服务": (1330, 300), "Supabase": (1330, 570)}
    for name, (x, y) in actors.items():
        d.ellipse((x - 22, y - 60, x + 22, y - 16), outline="#222222", width=3)
        d.line((x, y - 16, x, y + 70), fill="#222222", width=3)
        d.line((x - 45, y + 15, x + 45, y + 15), fill="#222222", width=3)
        d.line((x, y + 70, x - 35, y + 125), fill="#222222", width=3)
        d.line((x, y + 70, x + 35, y + 125), fill="#222222", width=3)
        d.text((x - 45, y + 135), name, font=pil_font(20), fill=pc(TEXT))

    cases = [
        ("登录系统", (420, 210, 650, 285)),
        ("新增原始素材", (410, 350, 690, 430)),
        ("AI 单条解析", (760, 350, 1030, 430)),
        ("一键体系化", (410, 500, 690, 580)),
        ("查看知识体系", (760, 500, 1040, 580)),
        ("知识问答回存", (570, 650, 870, 730)),
        ("导出个人数据", (900, 650, 1160, 730)),
    ]
    centers = {}
    for label, box in cases:
        d.ellipse(box, fill=pc("F8FBFE"), outline=pc(BLUE), width=3)
        x1, y1, x2, y2 = box
        centers[label] = ((x1 + x2) // 2, (y1 + y2) // 2)
        draw_wrapped(d, (x1 + 20, y1 + 25), label, pil_font(20), max_width=x2 - x1 - 40, center=True)

    for label in ["登录系统"]:
        d.line((155, 320, centers[label][0] - 115, centers[label][1]), fill="#333333", width=2)
    for label in ["新增原始素材", "一键体系化", "查看知识体系", "知识问答回存", "导出个人数据"]:
        d.line((155, 580, centers[label][0] - 135, centers[label][1]), fill="#333333", width=2)
    for label in ["AI 单条解析", "一键体系化", "知识问答回存"]:
        d.line((1285, 360, centers[label][0] + 135, centers[label][1]), fill="#666666", width=2)
    for label in ["登录系统", "新增原始素材", "查看知识体系", "导出个人数据"]:
        d.line((1285, 630, centers[label][0] + 120, centers[label][1]), fill="#666666", width=2)
    return save_image("fig_usecase.png", img)


def make_architecture() -> Path:
    img = Image.new("RGB", (1500, 900), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    d.text((560, 35), "序知系统总体架构图", font=pil_font(28), fill=pc(BLUE))
    boxes = {
        "用户界面层\nNext.js App Router\nWorkspace / Materials / Knowledge / QA / Settings": (120, 140, 520, 280),
        "业务接口层\nNext.js API Routes\n鉴权、参数校验、业务编排": (570, 140, 970, 280),
        "AI 能力层\nOpenAI 兼容接口\n解析、体系化、问答、OCR/转写": (1020, 140, 1400, 280),
        "认证与权限\nSupabase Auth\nCookie Session / RLS": (180, 420, 520, 545),
        "业务数据层\nPostgreSQL + pgvector\n素材、解析、节点、关系、版本": (580, 420, 980, 545),
        "文件与导出\nStorage / JSON / Markdown\n附件、导入、导出": (1040, 420, 1380, 545),
        "质量保障\nTypeScript / Build / RLS\n类型检查、构建验证、用户隔离": (500, 665, 1000, 785),
    }
    for label, box in boxes.items():
        draw_box(d, box, label, fill=pc("F8FBFE"), outline=pc(BLUE))
    arrow(d, (520, 210), (570, 210), fill=pc(BLUE))
    arrow(d, (970, 210), (1020, 210), fill=pc(BLUE))
    arrow(d, (770, 280), (770, 420), fill=pc(BLUE))
    arrow(d, (600, 260), (430, 420), fill="#666666")
    arrow(d, (870, 260), (1110, 420), fill="#666666")
    arrow(d, (780, 545), (750, 665), fill="#666666")
    return save_image("fig_architecture.png", img)


def make_er() -> Path:
    img = Image.new("RGB", (1600, 1100), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    d.text((635, 35), "序知核心数据模型图", font=pil_font(28), fill=pc(BLUE))
    entities = {
        "materials\nid PK\nuser_id FK\nraw_content\nstatus\nsource_type": (80, 150, 390, 310),
        "material_analysis\nid PK\nmaterial_id FK\ncore_meaning\nknowledge_type\nkeywords": (80, 430, 390, 610),
        "material_chunks\nid PK\nmaterial_id FK\nchunk_index\ncontent\nembedding vector": (80, 735, 390, 925),
        "knowledge_topics\nid PK\nparent_id FK\ntitle\nlevel\nsort_order": (650, 150, 980, 330),
        "knowledge_nodes\nid PK\ntopic_id FK\ntitle\ncontent\nnode_type": (650, 440, 980, 630),
        "node_material_links\nid PK\nnode_id FK\nmaterial_id FK\nrelevance_score": (650, 760, 980, 925),
        "knowledge_edges\nid PK\nsource_node_id FK\ntarget_node_id FK\nedge_type\nconfidence": (1170, 440, 1500, 645),
        "knowledge_versions\nid PK\njob_id FK\nversion_number\nsnapshot jsonb": (1170, 150, 1500, 330),
        "reconstruction_jobs\nid PK\nstatus\ninput_material_ids\nstarted_at\nfinished_at": (1170, 760, 1500, 940),
    }
    centers = {}
    for label, box in entities.items():
        draw_box(d, box, label, fill=pc("F8FBFE"), outline=pc(BLUE))
        x1, y1, x2, y2 = box
        centers[label.split("\n")[0]] = ((x1 + x2) // 2, (y1 + y2) // 2)
    connections = [
        ("materials", "material_analysis", "1 : 0..1"),
        ("materials", "material_chunks", "1 : N"),
        ("materials", "node_material_links", "1 : N"),
        ("knowledge_topics", "knowledge_nodes", "1 : N"),
        ("knowledge_nodes", "node_material_links", "1 : N"),
        ("knowledge_nodes", "knowledge_edges", "1 : N"),
        ("reconstruction_jobs", "knowledge_versions", "1 : N"),
        ("knowledge_versions", "knowledge_topics", "1 : N"),
    ]
    for a, b, label in connections:
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        d.line((x1, y1, x2, y2), fill="#555555", width=3)
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        d.rectangle((mx - 45, my - 16, mx + 45, my + 16), fill=pc("FFFFFF"))
        d.text((mx - 35, my - 12), label, font=pil_font(15), fill="#333333")
    return save_image("fig_er.png", img)


def make_sequence(name: str, title: str, steps: list[tuple[str, str, str]], participants: list[str]) -> Path:
    w, h = 1600, 980
    img = Image.new("RGB", (w, h), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    d.text((w // 2 - 180, 35), title, font=pil_font(28), fill=pc(BLUE))
    xs = [150 + i * ((w - 300) // (len(participants) - 1)) for i in range(len(participants))]
    y_top, y_bottom = 120, 880
    for x, p in zip(xs, participants):
        d.rectangle((x - 90, y_top, x + 90, y_top + 55), fill=pc("F8FBFE"), outline=pc(BLUE), width=2)
        draw_wrapped(d, (x - 80, y_top + 16), p, pil_font(17), max_width=160, center=True)
        d.line((x, y_top + 55, x, y_bottom), fill="#888888", width=2)
    y = 220
    p_index = {p: i for i, p in enumerate(participants)}
    for src, dst, label in steps:
        x1, x2 = xs[p_index[src]], xs[p_index[dst]]
        arrow(d, (x1, y), (x2, y), fill="#333333", width=2)
        tx = min(x1, x2) + 20
        d.rectangle((tx, y - 30, tx + 310, y - 3), fill=pc("FFFFFF"))
        d.text((tx + 6, y - 29), label, font=pil_font(16), fill=pc(TEXT))
        y += 72
    return save_image(name, img)


def make_flow(name: str, title: str, nodes: list[str]) -> Path:
    img = Image.new("RGB", (1500, 900), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    d.text((560, 35), title, font=pil_font(28), fill=pc(BLUE))
    y = 130
    for idx, node in enumerate(nodes):
        x1, x2 = (250, 1250)
        fill = pc("F8FBFE") if idx % 2 == 0 else pc("FFFFFF")
        d.rounded_rectangle((x1, y, x2, y + 70), radius=10, fill=fill, outline=pc(BLUE), width=2)
        d.text((x1 + 25, y + 20), f"{idx + 1}. {node}", font=pil_font(20), fill=pc(TEXT))
        if idx < len(nodes) - 1:
            arrow(d, (750, y + 70), (750, y + 115), fill=pc(BLUE))
        y += 115
    return save_image(name, img)


def make_wireframe() -> Path:
    img = Image.new("RGB", (1500, 900), pc("FFFFFF"))
    d = ImageDraw.Draw(img)
    d.text((560, 35), "工作台三栏界面示意图", font=pil_font(28), fill=pc(BLUE))
    d.rectangle((90, 115, 1410, 805), outline=pc(BLUE), width=3)
    d.rectangle((90, 115, 1410, 180), fill=pc(LIGHT_BLUE), outline=pc(BLUE), width=2)
    d.text((120, 135), "序知 Workspace：新增素材  |  一键体系化  |  原始视图  |  体系视图  |  重构历史", font=pil_font(22), fill=pc(TEXT))
    d.rectangle((110, 210, 455, 770), outline=pc(GRID_BLUE), width=2)
    d.rectangle((480, 210, 970, 770), outline=pc(GRID_BLUE), width=2)
    d.rectangle((995, 210, 1390, 770), outline=pc(GRID_BLUE), width=2)
    d.text((210, 230), "原始素材列表", font=pil_font(22), fill=pc(BLUE))
    d.text((625, 230), "素材输入/详情", font=pil_font(22), fill=pc(BLUE))
    d.text((1110, 230), "AI 知识体系树", font=pil_font(22), fill=pc(BLUE))
    for i in range(5):
        d.rounded_rectangle((135, 290 + i * 85, 430, 350 + i * 85), radius=6, fill=pc("F8FBFE"), outline=pc(GRID_BLUE))
        d.text((155, 307 + i * 85), f"素材 {i + 1}  pending/analyzed", font=pil_font(16), fill=pc(TEXT))
    d.rounded_rectangle((520, 300, 930, 520), radius=6, fill=pc("FFFFFF"), outline=pc(GRID_BLUE))
    draw_wrapped(d, (545, 325), "用户可以直接粘贴杂乱内容；系统保留原始素材，并在解析后显示核心含义、关键词和来源引用。", pil_font(18), max_width=360)
    d.rounded_rectangle((535, 570, 915, 640), radius=6, fill=pc(LIGHT_BLUE), outline=pc(BLUE))
    d.text((650, 592), "保存并解析", font=pil_font(22), fill=pc(TEXT))
    tree = ["AI 个人知识库", "  ├─ 素材入库", "  ├─ AI 解析", "  └─ 体系化重构", "      ├─ 节点", "      └─ 关联"]
    for i, line in enumerate(tree):
        d.text((1035, 300 + i * 55), line, font=pil_font(20), fill=pc(TEXT))
    return save_image("fig_workspace.png", img)


def make_all_figures() -> dict[str, Path]:
    return {
        "usecase": make_usecase(),
        "architecture": make_architecture(),
        "er": make_er(),
        "seq_material": make_sequence(
            "fig_seq_material.png",
            "新增素材并解析时序图",
            [
                ("用户", "前端工作台", "输入原始素材"),
                ("前端工作台", "API Routes", "POST /api/materials"),
                ("API Routes", "Supabase", "保存 raw_content"),
                ("前端工作台", "API Routes", "POST /api/analyze"),
                ("API Routes", "AI 服务", "请求结构化解析"),
                ("AI 服务", "API Routes", "返回解析 JSON"),
                ("API Routes", "Supabase", "写入 analysis/chunks"),
                ("API Routes", "前端工作台", "返回 analyzed 状态"),
            ],
            ["用户", "前端工作台", "API Routes", "AI 服务", "Supabase"],
        ),
        "seq_system": make_sequence(
            "fig_seq_systematize.png",
            "一键体系化重构时序图",
            [
                ("用户", "工作台", "点击一键体系化"),
                ("工作台", "系统化 API", "POST /api/systematize"),
                ("系统化 API", "Supabase", "读取 analyzed 素材"),
                ("系统化 API", "AI 服务", "生成主题树 JSON"),
                ("系统化 API", "Supabase", "重建 topics/nodes"),
                ("系统化 API", "Supabase", "写入 edges/version"),
                ("系统化 API", "工作台", "返回版本号"),
            ],
            ["用户", "工作台", "系统化 API", "AI 服务", "Supabase"],
        ),
        "flow_ingest": make_flow(
            "fig_flow_ingest.png",
            "素材入库处理流程图",
            ["用户自由输入原始素材", "写入 materials 并保留 raw_content", "内容分块生成 material_chunks", "调用 AI 解析核心含义和关键词", "更新素材状态并展示解析结果"],
        ),
        "flow_test": make_flow(
            "fig_flow_test.png",
            "系统测试流程图",
            ["准备测试环境和环境变量", "执行 TypeScript 类型检查", "执行 Next.js 生产构建", "验证核心 API 路由和页面路由", "整理测试结论和缺陷处理记录"],
        ),
        "workspace": make_wireframe(),
    }


def add_cover(doc: Document):
    add_para(doc, "软件工程实验报告", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26, color=BLUE)
    add_para(doc, "自选题目结合 AI 技术规范化流程开发软件", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, color=MUTED)
    add_para(doc, "", size=12)
    add_para(doc, "题目：序知 AI 个人体系化知识库", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, color=TEXT)
    add_para(doc, "输入无序碎片，输出终身知识体系", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=BLUE)
    add_para(doc, "", size=12)
    rows = [
        ["课程名称", "软件工程实践"],
        ["项目名称", "序知 AI 个人体系化知识库"],
        ["开发方式", "个人项目，结合 AI 辅助完成需求、设计、编码与测试"],
        ["技术路线", "Next.js + React + TypeScript + Supabase + pgvector + AI Provider"],
        ["报告日期", "2026 年 6 月 15 日"],
        ["姓名/学号", "（请在此处填写）"],
        ["指导教师", "（请在此处填写）"],
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, [4.2, 11.8])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=120, bottom=120)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cells[0].paragraphs[0].add_run(label)
        set_run_font(r1, size=11, color=TEXT, bold=True)
        r2 = cells[1].paragraphs[0].add_run(value)
        set_run_font(r2, size=11, color=TEXT)
    add_para(doc, "", size=12)
    add_note(doc, "报告说明", "本文档按照示例文档的四类软件工程材料组织，将需求获取、需求分析、概要设计、详细设计、测试与用户说明整合为一份总实验报告。")
    doc.add_page_break()


def add_toc(doc: Document):
    doc.add_heading("目录", level=1)
    items = [
        "第一部分 需求获取文档",
        "第二部分 需求分析文档",
        "第三部分 概要设计文档",
        "第四部分 详细设计与实现说明",
        "第五部分 测试与用户使用说明",
        "附录 项目运行与交付信息",
    ]
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=11, color=TEXT)
    doc.add_page_break()


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def build_report() -> None:
    ensure_dirs()
    figures = make_all_figures()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("摘要", level=1)
    add_para(doc, "序知是一款面向个人长期学习和经验沉淀的 AI 个人体系化知识库软件。项目核心不是让用户手工整理笔记，而是允许用户自由输入无序素材，由 AI 完成理解、归类、去重、排序、重构和知识网络维护。系统采用 Next.js、React、TypeScript、Supabase、PostgreSQL、pgvector 和 OpenAI 兼容接口构建，已实现素材入库、AI 单条解析、一键体系化、知识树展示、知识问答、知识关联、版本历史、数据导出、OCR/文档/音频入口等功能。")
    add_para(doc, "本报告按照软件工程规范描述项目从需求获取、需求分析、概要设计、详细设计到测试与使用说明的全过程，并配套用例图、时序图、系统架构图、数据模型图、流程图和测试表格。")

    doc.add_heading("第一部分 需求获取文档", level=1)
    doc.add_heading("1.1 项目背景", level=2)
    add_para(doc, "普通用户在学习、工作和生活中会持续产生大量碎片化知识，例如课程笔记、会议记录、读书摘录、灵感想法、网页资料、问题复盘和经验总结。传统笔记软件通常要求用户自己建立分类、标题、层级和标签，长期使用后往往出现素材堆积、重复混乱、难以复用的问题。序知面向这一痛点，将产品目标聚焦为“用户自由无序录入，AI 全权体系化重构”。")
    add_para(doc, "项目借鉴 Karpathy LLM Knowledge Base 中“raw 素材层、AI 维护的知识网络层、规则手册层”的思想，但在工程实现上采用数据库结构化知识网络与向量检索结合的方式，保证多用户、版本、权限、导出和长期扩展能力。")

    doc.add_heading("1.2 产品目标", level=2)
    add_table(doc, "表 1 产品目标表", ["编号", "目标", "说明"], [
        ["G1", "降低整理负担", "用户只负责输入原始素材，不需要手动分类、排版和建目录。"],
        ["G2", "形成个人知识体系", "AI 将零散素材重构为主题、分支、知识节点和节点关系。"],
        ["G3", "保留事实来源", "所有 AI 知识节点必须能追溯到原始素材，避免脱离用户内容编造。"],
        ["G4", "持续演化", "通过版本历史、知识关联和健康检查，让知识库随素材增长而演化。"],
    ], [1.4, 4.0, 10.8])

    doc.add_heading("1.3 目标用户与使用场景", level=2)
    add_table(doc, "表 2 目标用户与场景表", ["用户类型", "典型场景", "核心需求"], [
        ["学生", "网课、教材、论文和复习笔记混合积累", "自动整理知识点，形成课程体系和复习路径。"],
        ["职场用户", "会议纪要、项目复盘、方案资料长期沉淀", "把零散经验转为可复用的方法库和项目知识库。"],
        ["创作者", "灵感、素材、摘录和观点持续积累", "自动聚合专题，减少素材丢失和重复整理。"],
        ["终身学习者", "多领域资料长期输入", "构建跨主题关联，形成个人知识网络。"],
    ], [2.6, 6.0, 7.6])

    doc.add_heading("1.4 功能需求", level=2)
    add_table(doc, "表 3 功能需求表", ["编号", "功能", "优先级", "需求描述"], [
        ["FR-01", "素材自由入库", "高", "支持用户录入文本、粘贴资料、OCR 文本、音频转写和文档解析结果，并保留原始内容。"],
        ["FR-02", "AI 单条解析", "高", "对每条素材提取核心含义、有效信息、冗余信息、主题、关键词和知识类型。"],
        ["FR-03", "一键体系化", "高", "读取已解析素材，生成一级主题、二级分支、知识节点和来源引用。"],
        ["FR-04", "双视图展示", "高", "支持原始素材视图和 AI 体系化视图切换。"],
        ["FR-05", "知识网络", "中", "建立知识节点之间的 related 等关系，并在知识详情中展示关联。"],
        ["FR-06", "知识问答", "中", "基于用户知识库内容回答问题，避免脱离素材生成。"],
        ["FR-07", "版本历史", "中", "每次体系化生成版本快照，支持历史查看和差异比较。"],
        ["FR-08", "数据导出", "中", "支持导出用户数据 JSON 和 Markdown。"],
        ["FR-09", "设置管理", "中", "支持模型选择、API Key 配置和隐私说明。"],
    ], [1.5, 3.0, 1.8, 9.9])

    doc.add_heading("1.5 非功能需求", level=2)
    add_table(doc, "表 4 非功能需求表", ["类别", "要求", "实现方式"], [
        ["安全性", "用户只能访问自己的数据", "Supabase Auth + RLS + 服务端 user_id 过滤。"],
        ["可维护性", "前后端类型一致", "TypeScript 类型定义和结构化 AI JSON 输出。"],
        ["可扩展性", "后续可增加多模态素材和增量重构", "素材层、解析层、知识网络层分离。"],
        ["可靠性", "AI 失败时不能破坏原始素材", "原始素材独立保存，解析和重构状态可失败回滚。"],
        ["可用性", "初次使用者可快速理解", "三栏工作台、状态徽章、知识树和设置页。"],
    ], [2.6, 5.2, 8.4])

    doc.add_heading("第二部分 需求分析文档", level=1)
    doc.add_heading("2.1 执行者识别", level=2)
    add_table(doc, "表 5 执行者说明表", ["执行者", "职责"], [
        ["访客", "访问登录页，通过 Magic Link 登录系统。"],
        ["登录用户", "录入素材、查看解析、触发体系化、问答、导出数据。"],
        ["AI 服务", "完成单条解析、体系化重构、问答生成和内容识别。"],
        ["Supabase", "提供身份认证、数据库、存储、RLS 权限和 pgvector 能力。"],
    ], [3.5, 12.5])
    add_image(doc, figures["usecase"], "图 1 序知系统用例图")

    doc.add_heading("2.2 核心用例说明", level=2)
    add_table(doc, "表 6 核心用例规约表", ["用例", "前置条件", "主流程", "后置结果"], [
        ["新增素材", "用户已登录", "输入标题和原始内容，提交到 /api/materials", "materials 表新增记录，状态为 pending。"],
        ["AI 解析素材", "素材存在且用户有权限", "调用 /api/analyze，AI 返回结构化 JSON", "写入 material_analysis 和 material_chunks，状态更新为 analyzed。"],
        ["一键体系化", "存在 analyzed 素材", "调用 /api/systematize，AI 生成主题树", "写入 topics、nodes、links、edges 和版本快照。"],
        ["知识问答", "用户已登录且有素材/节点", "提交问题，服务端检索用户知识上下文", "返回基于知识库的回答，可作为后续回存素材。"],
        ["导出数据", "用户已登录", "调用导出接口生成 JSON/Markdown", "用户获得自己的知识库备份。"],
    ], [2.6, 3.8, 6.1, 4.1])

    doc.add_heading("2.3 业务流程分析", level=2)
    add_image(doc, figures["flow_ingest"], "图 2 素材入库处理流程图")
    add_image(doc, figures["seq_material"], "图 3 新增素材并解析时序图")
    add_image(doc, figures["seq_system"], "图 4 一键体系化重构时序图")

    doc.add_heading("第三部分 概要设计文档", level=1)
    doc.add_heading("3.1 总体架构", level=2)
    add_para(doc, "序知采用前后端一体化架构。前端页面由 Next.js App Router 组织，业务接口由 Next.js API Routes 提供，数据层使用 Supabase PostgreSQL，身份认证使用 Supabase Auth，语义检索能力基于 pgvector，AI 能力通过 OpenAI 兼容接口封装。")
    add_image(doc, figures["architecture"], "图 5 序知系统总体架构图")

    doc.add_heading("3.2 模块划分", level=2)
    add_table(doc, "表 7 系统模块划分表", ["模块", "主要文件/目录", "功能说明"], [
        ["认证模块", "src/app/(auth)/login、src/lib/supabase", "Magic Link 登录、会话读取、路由保护。"],
        ["素材模块", "src/components/materials、/api/materials", "新增、编辑、删除、搜索、状态筛选和详情展示。"],
        ["AI 解析模块", "src/lib/ai/analyze.ts、/api/analyze", "调用模型解析素材，输出结构化 JSON。"],
        ["体系化模块", "src/lib/ai/systematize.ts、/api/systematize", "全库重构、落库主题/节点/引用/关系/版本。"],
        ["知识展示模块", "src/components/knowledge、/api/knowledge", "树状知识体系、节点详情、来源素材和知识关联。"],
        ["问答模块", "src/components/qa、/api/qa", "基于用户知识库上下文生成回答。"],
        ["设置与导出", "src/components/settings、/api/export", "模型设置、API Key、JSON/Markdown 导出和隐私说明。"],
    ], [3.0, 5.0, 8.0])

    doc.add_heading("3.3 数据库设计", level=2)
    add_para(doc, "数据库围绕三层知识结构设计：原始素材层负责保存事实来源，AI 理解层保存单条解析结果和向量分块，知识网络层保存主题、节点、来源引用、节点关系和体系化版本。所有业务表均通过 user_id 与 Supabase Auth 用户绑定，并启用行级安全策略。")
    add_image(doc, figures["er"], "图 6 序知核心数据模型图", width_cm=16.0)
    doc.add_page_break()
    add_table(doc, "表 8 核心数据表说明", ["表名", "作用", "关键字段"], [
        ["materials", "原始素材事实源", "raw_content、source_type、status、user_id"],
        ["material_analysis", "AI 单条解析结果", "core_meaning、useful_points、topics、keywords"],
        ["material_chunks", "长文本分块和向量", "content、chunk_index、embedding vector(1024)"],
        ["knowledge_topics", "一级主题和二级分支", "parent_id、level、sort_order、version_id"],
        ["knowledge_nodes", "体系化知识节点", "topic_id、title、content、node_type"],
        ["node_material_links", "节点到素材的证据引用", "node_id、material_id、relevance_score"],
        ["knowledge_edges", "节点之间的知识关系", "source_node_id、target_node_id、edge_type、confidence"],
        ["reconstruction_jobs", "体系化任务记录", "status、scope、input_material_ids、error_message"],
        ["knowledge_versions", "知识体系版本快照", "version_number、snapshot、summary"],
    ], [3.3, 5.0, 7.7])

    doc.add_heading("3.4 接口设计", level=2)
    add_table(doc, "表 9 主要接口设计表", ["接口", "方法", "功能", "权限"], [
        ["/api/materials", "GET/POST", "查询和新增素材", "登录用户"],
        ["/api/materials/[id]", "GET/PATCH/DELETE", "素材详情、编辑和删除", "素材所有者"],
        ["/api/analyze", "POST", "AI 单条解析", "登录用户"],
        ["/api/systematize", "POST/GET", "触发体系化和查询任务", "登录用户"],
        ["/api/knowledge", "GET", "获取知识主题树", "登录用户"],
        ["/api/knowledge/nodes/[id]", "GET/PATCH", "知识节点详情和更新", "登录用户"],
        ["/api/knowledge/edges/all", "GET", "获取知识图谱关系", "登录用户"],
        ["/api/qa", "POST", "知识库问答", "登录用户"],
        ["/api/export", "GET", "导出完整 JSON 数据", "登录用户"],
        ["/api/export/markdown", "GET", "导出 Markdown 知识体系", "登录用户"],
    ], [4.2, 2.2, 6.1, 3.5])

    doc.add_heading("第四部分 详细设计与实现说明", level=1)
    doc.add_heading("4.1 素材入库实现", level=2)
    add_para(doc, "素材入库是系统的事实入口。前端允许用户无约束输入标题和正文，后端将 raw_content 原样写入 materials 表。系统不会在原始素材层覆盖、整理或删除用户原文，所有 AI 解析结果都保存在独立表中。该设计保证后续体系化失败、AI 输出偏差或版本回滚时，用户的原始知识资产仍然完整。")
    add_table(doc, "表 10 素材状态流转表", ["状态", "含义", "触发条件"], [
        ["pending", "待解析", "用户新建素材后默认进入该状态。"],
        ["analyzing", "解析中", "用户触发解析或系统开始调用 AI。"],
        ["analyzed", "已解析", "AI 返回合法 JSON 并成功写入解析结果。"],
        ["failed", "解析失败", "AI 调用、JSON 解析或数据库写入失败。"],
    ], [3.0, 5.0, 8.0])

    doc.add_heading("4.2 AI 单条解析实现", level=2)
    add_para(doc, "AI 解析模块的输入是单条原始素材，输出为固定 JSON 结构，包括核心含义、有效信息、冗余信息、主题、知识类型、关键词和可能关联方向。服务端将输出保存到 material_analysis 表，并将长文本切分为 material_chunks，为后续语义检索和相似素材判断提供基础。")
    add_note(doc, "AI 边界", "解析结果必须基于用户提供的素材，不做无关扩写、不编造外部资料、不把普通问答能力扩展成娱乐聊天。")

    doc.add_heading("4.3 AI 体系化重构实现", level=2)
    add_para(doc, "体系化接口是序知的核心闭环。当前实现采用全量重建策略：读取当前用户所有 analyzed 素材及其解析结果，生成模型输入后调用 systematizeMaterials，得到结构化主题树。随后清空当前用户旧的 knowledge_edges、knowledge_nodes 和 knowledge_topics，按照一级主题、二级分支、知识节点、来源引用的顺序写入数据库，并生成版本快照。")
    add_table(doc, "表 11 体系化落库顺序表", ["步骤", "数据表", "说明"], [
        ["1", "reconstruction_jobs", "创建 running 任务，记录开始时间。"],
        ["2", "materials + material_analysis", "读取已解析素材，限制一次最多处理 100 条。"],
        ["3", "knowledge_topics", "写入一级主题和二级分支。"],
        ["4", "knowledge_nodes", "写入节点标题、内容、类型和排序。"],
        ["5", "node_material_links", "写入节点与来源素材引用关系。"],
        ["6", "knowledge_edges", "基于共同来源素材自动建立 related 关系。"],
        ["7", "knowledge_versions", "保存本次体系化 JSON 快照和摘要。"],
        ["8", "reconstruction_jobs", "更新任务为 completed 或 failed。"],
    ], [1.7, 4.2, 10.1])

    doc.add_heading("4.4 知识网络与 Karpathy 思路落地", level=2)
    add_para(doc, "序知并不是单次总结工具，而是持续演化的个人知识网络。原始素材层类似 raw，知识主题和节点层类似 AI 编译后的 wiki，规则约束则体现在 Prompt 和系统实现中。项目已经实现 knowledge_edges 表，用于表达节点之间的 related、prerequisite、supports、contradicts、extends、example_of、part_of、duplicate 等关系类型。当前 MVP 先基于共同来源素材自动生成 related 关系，后续可升级为 AI 判断语义关系和冲突关系。")

    doc.add_heading("4.5 前端交互设计", level=2)
    add_para(doc, "前端围绕高频工作流设计为三栏工作台：左侧是原始素材列表，中间是素材输入和详情区域，右侧是 AI 体系化知识树。用户不需要先思考分类，可以在同一界面完成素材录入、查看解析结果、触发体系化和阅读知识体系。")
    add_image(doc, figures["workspace"], "图 7 工作台三栏界面示意图")

    doc.add_heading("第五部分 测试与用户使用说明", level=1)
    doc.add_heading("5.1 测试环境", level=2)
    add_table(doc, "表 12 测试环境表", ["项目", "配置"], [
        ["操作系统", "Windows，本地开发目录 D:\\OrdKnow"],
        ["运行框架", "Next.js 16.2.7，React 19.2.7，TypeScript 6.0.3"],
        ["数据库", "Supabase PostgreSQL，启用 pgvector 扩展"],
        ["构建命令", "npm run lint；npm run build"],
        ["构建结果", "TypeScript 检查通过；生产构建通过，生成 27 个页面/接口路由。"],
    ], [4.0, 12.0])
    add_image(doc, figures["flow_test"], "图 8 系统测试流程图")

    doc.add_heading("5.2 测试用例", level=2)
    add_table(doc, "表 13 功能测试用例表", ["编号", "测试项", "操作步骤", "预期结果", "状态"], [
        ["TC-01", "类型检查", "执行 npm run lint", "tsc 无类型错误", "通过"],
        ["TC-02", "生产构建", "执行 npm run build", "Next.js 编译成功并输出路由表", "通过"],
        ["TC-03", "素材新增", "登录后在工作台输入素材并保存", "materials 表新增记录，前端列表刷新", "待联调"],
        ["TC-04", "AI 解析", "对 pending 素材触发解析", "状态更新为 analyzed 并显示解析结果", "待联调"],
        ["TC-05", "一键体系化", "点击体系化按钮", "生成主题、节点、引用和版本号", "待联调"],
        ["TC-06", "知识问答", "在问答页输入问题", "回答基于用户知识库上下文", "待联调"],
        ["TC-07", "数据导出", "在设置页点击导出", "下载用户完整数据 JSON 或 Markdown", "待联调"],
    ], [1.5, 2.8, 6.3, 4.0, 1.6])

    doc.add_heading("5.3 用户使用说明", level=2)
    add_table(doc, "表 14 用户操作流程表", ["步骤", "页面", "操作", "结果"], [
        ["1", "/login", "输入邮箱并通过 Magic Link 登录", "进入受保护主界面。"],
        ["2", "/workspace", "在中间面板输入任意原始素材", "素材以 raw_content 原样保存。"],
        ["3", "/materials", "查看素材列表和解析状态", "可按 pending、analyzed 等状态筛选。"],
        ["4", "/workspace", "点击一键体系化", "AI 将已解析素材重构为知识体系。"],
        ["5", "/knowledge", "浏览主题树和节点详情", "查看内容、来源素材和相关节点。"],
        ["6", "/qa", "围绕知识库提问", "获得基于个人知识库的回答。"],
        ["7", "/settings", "配置模型、API Key 或导出数据", "完成个性化设置和备份。"],
    ], [1.4, 3.0, 6.0, 5.6])

    doc.add_heading("5.4 缺陷与风险说明", level=2)
    add_table(doc, "表 15 风险与改进表", ["风险", "影响", "当前处理", "后续优化"], [
        ["AI 输出不稳定", "可能导致 JSON 格式错误或节点质量不一", "结构化 Prompt、错误捕获、状态标记 failed", "增加 JSON schema 校验和重试策略。"],
        ["全量重建成本", "素材量大时成本和耗时增加", "MVP 限制一次最多 100 条素材", "改为增量更新和局部重构。"],
        ["环境变量缺失", "登录或 AI 接口无法正常使用", "提供 .env.local.example", "增加设置向导和环境检查页面。"],
        ["知识关系较简单", "当前主要基于共同来源生成 related", "已建立 knowledge_edges 数据结构", "引入 AI 语义关系判断和冲突检测。"],
    ], [3.2, 4.2, 4.4, 4.2])

    doc.add_heading("附录 项目运行与交付信息", level=1)
    add_table(doc, "表 16 项目文件结构说明", ["路径", "说明"], [
        ["src/app", "Next.js 页面、布局和 API Routes。"],
        ["src/components", "素材、知识、问答、工作台、设置和基础 UI 组件。"],
        ["src/lib/ai", "AI 请求、Prompt、解析和体系化逻辑。"],
        ["src/lib/supabase", "Supabase 客户端、服务端和中间件封装。"],
        ["supabase/migrations", "数据库扩展、业务表、RLS、索引和触发器。"],
        ["src/types/index.ts", "核心业务类型定义。"],
        ["tools", "文档和样例生成脚本。"],
    ], [5.0, 11.0])
    add_para(doc, "运行步骤：复制 .env.local.example 为 .env.local，配置 Supabase URL、匿名 Key、AI Provider Key 和 Embedding Key；在 Supabase SQL Editor 中按顺序执行 migrations；执行 npm run dev 启动开发服务；访问 /login 完成登录后开始使用。")
    add_para(doc, "交付物包括：项目源代码、Supabase 数据库迁移、PRD 与技术路线文档、软件工程实验报告 DOCX。")

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
    print(REPORT_PATH)
