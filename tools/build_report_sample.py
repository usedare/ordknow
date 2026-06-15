from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_samples"
IMG_DIR = OUT_DIR / "images"
OUT_DOCX = OUT_DIR / "序知_图表与版式样例.docx"

BLUE = "4472C4"
LIGHT_BLUE = "D9E2F3"
BORDER_BLUE = "8EAADB"
TEXT = (20, 20, 20)
LINE = (30, 30, 30)
GRID = (232, 236, 242)

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_EN = "Times New Roman"

PIL_FONT_REG = r"C:\Windows\Fonts\simsun.ttc"
PIL_FONT_BOLD = r"C:\Windows\Fonts\simhei.ttf"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(PIL_FONT_BOLD if bold else PIL_FONT_REG, size=size)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def draw_grid(draw: ImageDraw.ImageDraw, w: int, h: int, step: int = 20):
    for x in range(0, w, step):
        draw.line([(x, 0), (x, h)], fill=GRID, width=1)
    for y in range(0, h, step):
        draw.line([(0, y), (w, y)], fill=GRID, width=1)


def draw_center_text(draw, box, text, fnt, fill=TEXT):
    x1, y1, x2, y2 = box
    tw, th = text_size(draw, text, fnt)
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2 - 2), text, font=fnt, fill=fill)


def draw_arrow(draw, start, end, fill=LINE, dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segments = 18
        for i in range(segments):
            if i % 2 == 0:
                a = i / segments
                b = (i + 1) / segments
                draw.line([(x1 + (x2 - x1) * a, y1 + (y2 - y1) * a), (x1 + (x2 - x1) * b, y1 + (y2 - y1) * b)], fill=fill, width=2)
    else:
        draw.line([start, end], fill=fill, width=2)
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 10
    p1 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
    p2 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def draw_actor(draw, x, y, label):
    draw.ellipse((x - 10, y, x + 10, y + 20), outline=LINE, width=2)
    draw.line((x, y + 20, x, y + 58), fill=LINE, width=2)
    draw.line((x - 22, y + 34, x + 22, y + 34), fill=LINE, width=2)
    draw.line((x, y + 58, x - 18, y + 86), fill=LINE, width=2)
    draw.line((x, y + 58, x + 18, y + 86), fill=LINE, width=2)
    tw, _ = text_size(draw, label, font(16))
    draw.text((x - tw / 2, y + 92), label, font=font(16), fill=TEXT)


def draw_usecase(draw, cx, cy, text):
    rx, ry = 80, 26
    draw.ellipse((cx - rx, cy - ry, cx + rx, cy + ry), outline=LINE, width=2, fill=(255, 255, 255))
    draw_center_text(draw, (cx - rx, cy - ry, cx + rx, cy + ry), text, font(15, True))


def make_usecase(path: Path):
    img = Image.new("RGB", (1000, 680), "white")
    draw = ImageDraw.Draw(img)
    draw_grid(draw, 1000, 680, 18)
    draw.rectangle((130, 60, 860, 610), outline=LINE, width=2)
    draw_center_text(draw, (430, 62, 560, 92), "序知平台", font(16, True))

    draw_actor(draw, 70, 245, "用户")
    draw_actor(draw, 930, 245, "AI")
    draw_actor(draw, 930, 475, "数据库")

    cases = [
        (300, 120, "新增素材"),
        (530, 120, "图片/文件导入"),
        (300, 220, "AI 单条解析"),
        (530, 220, "生成向量"),
        (300, 320, "一键体系化"),
        (530, 320, "知识网络维护"),
        (300, 420, "知识问答"),
        (530, 420, "问答回存"),
        (410, 535, "导出 Markdown / JSON"),
    ]
    for cx, cy, t in cases:
        draw_usecase(draw, cx, cy, t)

    for _, cy, _ in cases[:7]:
        draw.line((100, 290, 220, cy), fill=LINE, width=2)
    # 右侧参与者只连接与其直接相关的用例，避免关系线过密影响阅读。
    for cx, cy, _ in [cases[2], cases[4], cases[6], cases[7]]:
        draw.line((900, 290, cx + 80, cy), fill=LINE, width=1)
    for cx, cy, _ in [cases[0], cases[3], cases[5], cases[8]]:
        draw.line((900, 520, cx + 80, cy), fill=LINE, width=1)

    # include / extend 标注，保持示例 PDF 的 UML 视觉风格。
    draw_arrow(draw, (380, 220), (450, 220), dashed=True)
    draw.text((395, 195), "<<include>>", font=font(12), fill=TEXT)
    draw_arrow(draw, (380, 320), (450, 320), dashed=True)
    draw.text((395, 295), "<<include>>", font=font(12), fill=TEXT)
    draw_arrow(draw, (380, 420), (450, 420), dashed=True)
    draw.text((395, 395), "<<extend>>", font=font(12), fill=TEXT)

    img.save(path)


def make_sequence(path: Path):
    img = Image.new("RGB", (1000, 700), "white")
    draw = ImageDraw.Draw(img)
    f = font(15)
    fb = font(15, True)
    xs = [110, 300, 500, 700, 880]
    labels = ["用户", "前端工作台", "API Routes", "AI 服务", "Supabase"]
    for x, label in zip(xs, labels):
        draw.rounded_rectangle((x - 70, 45, x + 70, 85), radius=5, outline=LINE, fill=(238, 241, 248), width=2)
        draw_center_text(draw, (x - 70, 45, x + 70, 85), label, fb)
        draw.line((x, 85, x, 640), fill=(150, 150, 150), width=1)
        for yy in range(95, 640, 16):
            draw.line((x, yy, x, yy + 8), fill=(185, 185, 185), width=1)

    messages = [
        (0, 1, 130, "1. 输入原始素材"),
        (1, 2, 190, "2. POST /api/materials"),
        (2, 4, 250, "3. 保存 raw_content"),
        (1, 2, 315, "4. POST /api/analyze"),
        (2, 3, 385, "5. 请求结构化解析"),
        (3, 2, 445, "6. 返回 JSON 解析结果", True),
        (2, 4, 505, "7. 写入 analysis / chunks"),
        (4, 1, 575, "8. 返回 analyzed 状态", True),
    ]
    for item in messages:
        src, dst, y, label, *rest = item
        dashed = bool(rest and rest[0])
        draw_arrow(draw, (xs[src], y), (xs[dst], y), dashed=dashed)
        tx = min(xs[src], xs[dst]) + 10
        draw.text((tx, y - 24), label, font=f, fill=TEXT)

    draw.rectangle((220, 340, 950, 620), outline=LINE, width=2)
    draw.rectangle((220, 340, 275, 363), outline=LINE, width=2, fill=(255, 255, 255))
    draw.text((232, 344), "alt", font=fb, fill=TEXT)
    draw.text((285, 345), "[AI 解析成功]", font=f, fill=TEXT)
    draw.line((220, 470, 950, 470), fill=LINE, width=1)
    draw.text((285, 477), "[AI 或 Embedding 失败]", font=f, fill=TEXT)
    draw.text((305, 530), "Embedding 失败不阻断基础解析，素材仍可进入理解层。", font=f, fill=TEXT)

    img.save(path)


def make_architecture(path: Path):
    img = Image.new("RGB", (1000, 620), "white")
    draw = ImageDraw.Draw(img)
    fb = font(16, True)
    f = font(14)
    groups = [
        (90, 60, 910, 150, "View：Next.js 页面与组件", ["Workspace", "Materials", "Knowledge", "QA", "Settings"]),
        (90, 210, 910, 310, "Controller：API Routes", ["materials", "analyze", "systematize", "qa", "export"]),
        (90, 370, 910, 470, "Services：AI 与导入处理", ["DeepSeek/MiMo", "SiliconFlow", "OCR", "PDF/Word", "Audio"]),
        (90, 520, 910, 590, "Data：Supabase", ["Auth/RLS", "PostgreSQL", "pgvector", "Storage"]),
    ]
    for x1, y1, x2, y2, title, items in groups:
        draw.rectangle((x1, y1, x2, y2), outline=(190, 170, 210), width=2, fill=(255, 253, 247))
        draw.text((x1 + 12, y1 + 8), title, font=fb, fill=TEXT)
        item_w = 130
        gap = 20
        total = len(items) * item_w + (len(items) - 1) * gap
        sx = x1 + (x2 - x1 - total) // 2
        for i, item in enumerate(items):
            bx = sx + i * (item_w + gap)
            draw.rounded_rectangle((bx, y1 + 42, bx + item_w, y1 + 78), radius=4, outline=(210, 185, 145), fill=(255, 248, 230), width=2)
            draw_center_text(draw, (bx, y1 + 42, bx + item_w, y1 + 78), item, f)

    for y in [150, 310, 470]:
        draw_arrow(draw, (500, y + 10), (500, y + 55))
    draw.text((520, 165), "HTTP / fetch", font=f, fill=(110, 80, 130))
    draw.text((520, 325), "模型调用 / 文件解析", font=f, fill=(110, 80, 130))
    draw.text((520, 482), "结构化存储 / 向量检索", font=f, fill=(110, 80, 130))
    img.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER_BLUE, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia=FONT_SONG, ascii_font=FONT_EN, size=None, bold=None, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, FONT_HEI, FONT_EN, 16, True, "00AEEF")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, FONT_HEI, FONT_EN, 13, True, None)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    else:
        set_run_font(run, FONT_HEI, FONT_EN, 11.5, True, None)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, FONT_EN, 12)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, FONT_EN, 10.5)


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, FONT_EN, 10.5)


def add_blue_table(doc, title, headers, rows, widths):
    add_table_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = headers[idx]
        set_cell_shading(cell, BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, FONT_HEI, FONT_EN, 10.5, True, "FFFFFF")
        cell.width = Cm(widths[idx])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            set_cell_shading(cells[c_idx], LIGHT_BLUE if r_idx % 2 == 0 else "FFFFFF")
            set_cell_border(cells[c_idx])
            set_cell_margins(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[c_idx].width = Cm(widths[c_idx])
            para = cells[c_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.25
            for run in para.runs:
                set_run_font(run, FONT_SONG if c_idx else FONT_HEI, FONT_EN, 10.5, c_idx == 0)
    return table


def create_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    usecase = IMG_DIR / "usecase.png"
    sequence = IMG_DIR / "sequence.png"
    arch = IMG_DIR / "architecture.png"
    make_usecase(usecase)
    make_sequence(sequence)
    make_architecture(arch)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.6)

    styles = doc.styles
    styles["Normal"].font.name = FONT_EN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    r = title.add_run("序知 AI 个人体系化知识库\n图表与版式样例")
    set_run_font(r, FONT_HEI, FONT_EN, 18, True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("用于确认实验报告正文、表格、UML 图和图题样式")
    set_run_font(r, FONT_SONG, FONT_EN, 11)

    add_heading(doc, "1. 软件需求样例", 1)
    add_body(doc, "本样例仅用于确认最终实验报告的版式和图表样式。正文采用正式软件工程实验报告风格，图表内容基于当前 OrdKnow（序知）项目代码与产品说明，不引入未实现的功能。")
    add_heading(doc, "1.1 功能需求", 2)
    add_blue_table(
        doc,
        "表 1 功能需求表示例",
        ["模块", "核心功能"],
        [
            ["素材入库", "- 支持文本、图片 OCR、PDF/Word 文档、音频转写和网页抓取。\n- 原始内容以 raw_content 保留，不被 AI 覆盖。"],
            ["AI 解析", "- 对单条素材生成核心含义、知识类型、主题、关键词。\n- 解析结果写入 material_analysis。"],
            ["体系化重构", "- 将已解析素材重构为一级主题、二级分支和知识节点。\n- 每个节点保留来源素材引用。"],
            ["知识网络", "- 根据共同来源素材建立 related 关系。\n- 支持树状视图和知识图谱视图。"],
            ["知识问答", "- 后端读取当前用户素材生成回答。\n- 用户可将优质问答回存为新素材。"],
        ],
        [3.1, 12.4],
    )

    add_heading(doc, "1.2 用户用例", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(usecase), width=Cm(15.8))
    add_caption(doc, "图 1 用户用例图样例")

    add_body(doc, "如图 1 所示，序知平台以普通用户为核心执行者，AI 服务与数据库作为被动参与者。用户主要完成素材录入、AI 解析、知识体系生成、知识问答和数据导出等操作。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "2. 概要设计样例", 1)
    add_heading(doc, "2.1 系统总体架构", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(arch), width=Cm(15.8))
    add_caption(doc, "图 2 系统架构图样例")
    add_body(doc, "系统采用 Next.js 页面层、API Routes 控制层、AI 与导入服务层、Supabase 数据层的分层结构。前端负责交互和配置，后端负责权限校验、AI 调用、数据库写入和导出。")

    add_heading(doc, "2.2 数据模型关系", 2)
    add_blue_table(
        doc,
        "表 2 核心数据表说明示例",
        ["数据表", "作用说明", "关键关系"],
        [
            ["materials", "保存用户原始素材，是系统事实源。", "一对多关联 material_chunks / material_analysis"],
            ["material_analysis", "保存 AI 对单条素材的结构化理解。", "每条记录关联一个 material"],
            ["knowledge_topics", "保存一级主题和二级分支。", "parent_id 支持层级结构"],
            ["knowledge_nodes", "保存体系化后的知识节点。", "挂载到二级 topic"],
            ["knowledge_edges", "保存知识节点之间的关系。", "source_node_id / target_node_id"],
        ],
        [3.8, 7.2, 4.5],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "3. 详细设计样例", 1)
    add_heading(doc, "3.1 新增素材并解析时序", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(sequence), width=Cm(15.8))
    add_caption(doc, "图 3 新增素材并解析时序图样例")
    add_body(doc, "图 3 展示了从用户输入素材到系统完成 AI 解析的主要交互过程。素材先入库为 pending 状态，随后后端并行执行结构化解析和向量生成。若向量生成失败，系统仍保留基础解析结果。")

    add_heading(doc, "3.2 用例文字说明样例", 2)
    items = [
        ("主要参与者", "用户"),
        ("被动参与者", "前端工作台、API Routes、AI 服务、Supabase 数据库"),
        ("前置条件", "用户已登录，且进入素材管理或工作台页面。"),
        ("后置条件", "素材状态变为 analyzed，并产生可供体系化重构使用的解析结果。"),
        ("基本交互序列", "1. 用户输入标题和原始内容。\n2. 系统保存素材并标记为 pending。\n3. 用户点击 AI 解析。\n4. 系统调用模型生成结构化 JSON。\n5. 系统保存解析结果并更新素材状态。"),
        ("扩展交互序列", "4a. 若 AI 服务不可用，系统标记 failed 并提示错误。\n4b. 若 embedding 配置缺失，系统跳过向量生成，但保留基础解析结果。"),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        r1 = p.add_run(f"○ {label}：")
        set_run_font(r1, FONT_SONG, FONT_EN, 12, False)
        r2 = p.add_run(value)
        set_run_font(r2, FONT_SONG, FONT_EN, 12)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    create_docx()
