from __future__ import annotations

import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
IMG = OUT / "template_figures"
DOCX = OUT / "序知_软件工程文档合订本.docx"
PDF = OUT / "序知_软件工程文档合订本.pdf"

BLUE = "00A2E8"
TABLE_BLUE = "4472C4"
LIGHT_ROW = "D9EAF7"
LIGHTER = "EEF6FB"
GRID = "A6C8E6"
INK = "111111"


def font_file(preferred: list[str]) -> str:
    for name in preferred:
        p = Path(r"C:\Windows\Fonts") / name
        if p.exists():
            return str(p)
    return str(Path(r"C:\Windows\Fonts\msyh.ttc"))


FONT_SONG = font_file(["simsun.ttc", "msyh.ttc"])
FONT_HEI = font_file(["simhei.ttf", "msyh.ttc"])
FONT_KAI = font_file(["simkai.ttf", "STKAITI.TTF", "simsun.ttc"])


def pf(size: int, kind: str = "song"):
    f = {"song": FONT_SONG, "hei": FONT_HEI, "kai": FONT_KAI}.get(kind, FONT_SONG)
    return ImageFont.truetype(f, size=size)


def pc(c: str) -> str:
    return c if c.startswith("#") else f"#{c}"


def set_run(run, font="宋体", size=10.5, color=INK, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def para(doc, text="", style=None, align=None, font="宋体", size=10.5, bold=False, color=INK, before=0, after=4):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, font=font, size=size, color=color, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level > 1 else 12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, font="黑体", size={1: 14, 2: 12, 3: 11}.get(level, 10.5), color=BLUE if level == 1 else INK, bold=True)
    return p


def bullet(doc, text, indent=0):
    p = para(doc, "", after=2)
    p.paragraph_format.left_indent = Cm(0.75 + indent * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run("○  " + text)
    set_run(r, size=10.5)


def numbered(doc, text, n):
    p = para(doc, "", after=2)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run(f"{n}.{text}")
    set_run(r, size=10.5)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=GRID):
    tc_pr = cell._tc.get_or_add_tcPr()
    b = tc_pr.first_child_found_in("w:tcBorders")
    if b is None:
        b = OxmlElement("w:tcBorders")
        tc_pr.append(b)
    for edge in ["top", "left", "bottom", "right"]:
        e = b.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), color)


def margins(cell, top=80, bottom=80, start=110, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    m = tc_pr.first_child_found_in("w:tcMar")
    if m is None:
        m = OxmlElement("w:tcMar")
        tc_pr.append(m)
    for key, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        e = m.find(qn(f"w:{key}"))
        if e is None:
            e = OxmlElement(f"w:{key}")
            m.append(e)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")


def _table_variant(headers, caption, variant):
    if variant:
        return variant
    caption = caption or ""
    first = headers[0] if headers else ""
    if first == "字段":
        return "data"
    if first == "编号" or "测试" in caption:
        return "test"
    if "截图" in caption or "证据" in caption:
        return "evidence"
    if len(headers) == 2 and first in {"项目", "说明项", "设计项", "部署项", "验证项", "检查点"}:
        return "kv"
    if len(headers) == 2:
        return "plain"
    return "blue"


def _style_tokens(variant):
    return {
        "blue": {
            "header_fill": TABLE_BLUE,
            "header_text": "FFFFFF",
            "row_fills": [LIGHT_ROW, "FFFFFF"],
            "border": GRID,
            "header_font": "黑体",
            "body_size": 9.5,
        },
        "plain": {
            "header_fill": "F3F6FA",
            "header_text": INK,
            "row_fills": ["FFFFFF", "F8FAFC"],
            "border": "CCD6E0",
            "header_font": "黑体",
            "body_size": 9.5,
        },
        "kv": {
            "header_fill": "FFFFFF",
            "header_text": "1F4E79",
            "row_fills": ["FFFFFF"],
            "label_fill": "EEF3F8",
            "border": "B8C7D6",
            "header_font": "黑体",
            "body_size": 9.5,
        },
        "data": {
            "header_fill": "3F4B5B",
            "header_text": "FFFFFF",
            "row_fills": ["F7F9FB", "FFFFFF"],
            "border": "B9C0C8",
            "header_font": "黑体",
            "body_size": 9.2,
        },
        "test": {
            "header_fill": "595959",
            "header_text": "FFFFFF",
            "row_fills": ["F2F2F2", "FFFFFF"],
            "border": "C8C8C8",
            "header_font": "黑体",
            "body_size": 9.0,
        },
        "evidence": {
            "header_fill": "1F4E79",
            "header_text": "FFFFFF",
            "row_fills": ["EAF3F8", "FFFFFF"],
            "border": "9EB6CC",
            "header_font": "黑体",
            "body_size": 9.3,
        },
    }[variant]


def table(doc, headers, rows, widths, caption=None, variant=None, trailing_space=True):
    variant = _table_variant(headers, caption, variant)
    tokens = _style_tokens(variant)
    if caption:
        p = para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体", size=10.5, bold=True, after=4)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, tokens["header_fill"])
        borders(c, tokens["border"])
        margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, font=tokens["header_font"], size=10, color=tokens["header_text"], bold=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.width = Cm(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if variant == "kv" and i == 0:
                shade(c, tokens["label_fill"])
            else:
                fills = tokens["row_fills"]
                shade(c, fills[ri % len(fills)])
            borders(c, tokens["border"])
            margins(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(val) < 10 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run(r, size=tokens["body_size"], color=INK, bold=(i == 0 and (variant == "kv" or len(val) < 8)))
    if trailing_space:
        para(doc, "", after=2)
    return t


def callout(doc, title, body, tone="info"):
    palette = {
        "info": ("EAF3F8", "1F4E79"),
        "warn": ("FFF4D6", "7A5A00"),
        "ok": ("EAF7EF", "276749"),
        "plain": ("F5F5F5", "555555"),
    }
    fill, accent = palette.get(tone, palette["info"])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Cm(15.4)
    shade(cell, fill)
    borders(cell, accent)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, font="黑体", size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.22
    r2 = p2.add_run(body)
    set_run(r2, size=9.8, color=INK)
    para(doc, "", after=2)
    return t


def configure(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.7)
    sec.right_margin = Cm(2.4)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("序知 AI 个人体系化知识库软件工程文档")
    set_run(r, size=8.5, color="666666")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def wrap(draw, text, font, max_width):
    lines, line = [], ""
    for ch in text:
        trial = line + ch
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not line:
            line = trial
        else:
            lines.append(line)
            line = ch
    if line:
        lines.append(line)
    return lines


def grid(draw, w, h, step=22):
    for x in range(0, w, step):
        draw.line((x, 0, x, h), fill="#EEF2F5", width=1)
    for y in range(0, h, step):
        draw.line((0, y, w, y), fill="#EEF2F5", width=1)


def centered(draw, box, text, font, fill="#111111"):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, font, x2 - x1 - 16)
    total = len(lines) * (font.size + 4)
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - bbox[2]) / 2, y), line, font=font, fill=fill)
        y += font.size + 4


def arrow(draw, x1, y1, x2, y2, dash=False):
    if dash:
        steps = 18
        for i in range(0, steps, 2):
            a = i / steps
            b = (i + 1) / steps
            draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill="#222222", width=2)
    else:
        draw.line((x1, y1, x2, y2), fill="#222222", width=2)
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 10, y2 - 5), (x2 - 10, y2 + 5)]
    else:
        pts = [(x2, y2), (x2 + 10, y2 - 5), (x2 + 10, y2 + 5)]
    draw.polygon(pts, fill="#222222")


def save(img, name):
    p = IMG / name
    img.save(p)
    return p


def cover_image():
    img = Image.new("RGB", (1240, 1754), "#F3F0EE")
    d = ImageDraw.Draw(img)
    d.ellipse((-200, 1350, 1050, 2060), fill="#DDE7E4")
    d.ellipse((430, 1280, 1540, 2020), fill="#E7E0D6")
    for i in range(7):
        d.arc((220 + i * 55, 520 + i * 80, 740 + i * 70, 820 + i * 90), 185, 345, fill="#E6C778", width=5)
    d.text((220, 135), "序知", font=pf(150, "kai"), fill="#0A0A0A")
    for idx, ch in enumerate("AI个人体系化知识库"):
        d.text((595, 335 + idx * 115), ch, font=pf(76, "kai"), fill="#0A0A0A")
    for idx, ch in enumerate("软件工程文档合订本"):
        d.text((815, 660 + idx * 62), ch, font=pf(44, "kai"), fill="#0A0A0A")
    d.text((115, 1250), "制作说明：", font=pf(42, "kai"), fill="#0A0A0A")
    for i, txt in enumerate(["项目名称：序知 AI 个人体系化知识库", "课程：软件工程", "指导老师：胡刚", "日期：2026 年 6 月"]):
        d.text((115, 1320 + i * 60), txt, font=pf(34, "kai"), fill="#0A0A0A")
    return save(img, "cover.png")


def usecase():
    img = Image.new("RGB", (1500, 1000), "white")
    d = ImageDraw.Draw(img)
    grid(d, 1500, 1000)
    d.rectangle((180, 90, 1250, 910), outline="#222222", width=2)
    d.text((210, 110), "序知 AI 个人体系化知识库", font=pf(24, "hei"), fill="#111111")
    actors = {"用户": (90, 390), "AI 服务": (1370, 355), "数据库": (1370, 660)}
    for name, (x, y) in actors.items():
        d.ellipse((x - 18, y - 55, x + 18, y - 19), outline="#222222", width=2)
        d.line((x, y - 19, x, y + 52), fill="#222222", width=2)
        d.line((x - 36, y + 8, x + 36, y + 8), fill="#222222", width=2)
        d.line((x, y + 52, x - 30, y + 100), fill="#222222", width=2)
        d.line((x, y + 52, x + 30, y + 100), fill="#222222", width=2)
        d.text((x - 36, y + 110), name, font=pf(21, "song"), fill="#111111")
    cases = [
        ("注册/登录", (380, 160, 570, 220)),
        ("录入原始素材", (340, 280, 610, 350)),
        ("AI 单条解析", (720, 280, 980, 350)),
        ("一键体系化", (340, 430, 610, 500)),
        ("查看知识体系", (720, 430, 1010, 500)),
        ("知识问答", (340, 590, 610, 660)),
        ("导出数据", (720, 590, 980, 660)),
        ("查看版本历史", (520, 750, 810, 820)),
    ]
    centers = {}
    for text, box in cases:
        d.ellipse(box, fill="#FFFFFF", outline="#222222", width=2)
        centered(d, box, text, pf(21, "song"))
        centers[text] = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
    for text in ["注册/登录", "录入原始素材", "一键体系化", "查看知识体系", "知识问答", "导出数据", "查看版本历史"]:
        x, y = centers[text]
        d.line((130, 445, x - 130, y), fill="#222222", width=2)
    for text in ["AI 单条解析", "一键体系化", "知识问答"]:
        x, y = centers[text]
        d.line((1328, 415, x + 130, y), fill="#222222", width=2)
    for text in ["注册/登录", "录入原始素材", "查看知识体系", "导出数据", "查看版本历史"]:
        x, y = centers[text]
        d.line((1328, 720, x + 130, y), fill="#222222", width=2)
    d.text((640, 260), "<<include>>", font=pf(18, "song"), fill="#555555")
    d.line((610, 315, 720, 315), fill="#555555", width=1)
    return save(img, "usecase.png")


def seq(name, title, participants, steps):
    img = Image.new("RGB", (1500, 950), "white")
    d = ImageDraw.Draw(img)
    xs = [170 + i * (1160 // (len(participants) - 1)) for i in range(len(participants))]
    d.text((95, 40), title, font=pf(26, "hei"), fill="#111111")
    top = 110
    for x, p in zip(xs, participants):
        d.rounded_rectangle((x - 85, top, x + 85, top + 46), radius=4, fill="#E9E8F4", outline="#777777")
        centered(d, (x - 85, top, x + 85, top + 46), p, pf(18, "song"))
        d.line((x, top + 46, x, 850), fill="#888888", width=1)
    y = 205
    idx = {p: i for i, p in enumerate(participants)}
    for s, t, label, ret in steps:
        x1, x2 = xs[idx[s]], xs[idx[t]]
        arrow(d, x1, y, x2, y, dash=ret)
        d.rectangle((min(x1, x2) + 18, y - 30, min(x1, x2) + 350, y - 6), fill="white")
        d.text((min(x1, x2) + 22, y - 30), label, font=pf(17, "song"), fill="#111111")
        y += 72
    return save(img, name)


def architecture():
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    grid(d, 1500, 900, 25)
    d.text((80, 40), "序知软件体系结构图", font=pf(28, "hei"), fill="#111111")
    layers = [
        ("表示层", ["登录页", "工作台", "素材页", "知识体系页", "问答页", "设置页"]),
        ("接口层", ["API Routes", "Server Actions", "Proxy 鉴权"]),
        ("业务层", ["素材管理", "AI 解析", "体系化重构", "知识网络", "问答检索", "导出"]),
        ("数据层", ["Supabase Auth", "PostgreSQL", "pgvector", "Storage"]),
        ("外部服务", ["AI Provider", "Embedding Provider", "OCR/Transcribe"]),
    ]
    y = 115
    for layer, items in layers:
        d.rectangle((130, y, 1370, y + 105), outline="#222222", width=2)
        d.rectangle((130, y, 260, y + 105), fill="#F1F1F7", outline="#222222", width=2)
        centered(d, (130, y, 260, y + 105), layer, pf(22, "hei"))
        x = 300
        for item in items:
            d.rounded_rectangle((x, y + 25, x + 150, y + 78), radius=5, fill="#FFFFFF", outline="#777777")
            centered(d, (x, y + 25, x + 150, y + 78), item, pf(17, "song"))
            x += 170
        if y < 580:
            arrow(d, 750, y + 105, 750, y + 145)
        y += 145
    return save(img, "architecture.png")


def er():
    img = Image.new("RGB", (1500, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 40), "序知核心数据表关系图", font=pf(28, "hei"), fill="#111111")
    entities = {
        "materials": (80, 120, 350, 245),
        "material_analysis": (80, 360, 350, 500),
        "material_chunks": (80, 620, 350, 760),
        "knowledge_topics": (610, 120, 910, 260),
        "knowledge_nodes": (610, 380, 910, 520),
        "node_material_links": (610, 650, 910, 790),
        "knowledge_edges": (1120, 380, 1420, 530),
        "knowledge_versions": (1120, 120, 1420, 260),
        "reconstruction_jobs": (1120, 650, 1420, 790),
    }
    fields = {
        "materials": ["id PK", "user_id FK", "raw_content", "status"],
        "material_analysis": ["id PK", "material_id FK", "core_meaning", "topics"],
        "material_chunks": ["id PK", "material_id FK", "content", "embedding"],
        "knowledge_topics": ["id PK", "parent_id FK", "title", "level"],
        "knowledge_nodes": ["id PK", "topic_id FK", "title", "content"],
        "node_material_links": ["id PK", "node_id FK", "material_id FK", "score"],
        "knowledge_edges": ["id PK", "source_node_id FK", "target_node_id FK", "edge_type"],
        "knowledge_versions": ["id PK", "job_id FK", "version_number", "snapshot"],
        "reconstruction_jobs": ["id PK", "status", "input_material_ids", "error_message"],
    }
    centers = {}
    for name, box in entities.items():
        d.rectangle(box, fill="#FFFFFF", outline="#222222", width=2)
        d.rectangle((box[0], box[1], box[2], box[1] + 34), fill="#F1F1F7", outline="#222222", width=1)
        centered(d, (box[0], box[1], box[2], box[1] + 34), name, pf(17, "hei"))
        yy = box[1] + 45
        for f in fields[name]:
            d.text((box[0] + 15, yy), f, font=pf(15, "song"), fill="#111111")
            yy += 22
        centers[name] = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
    for a, b, lab in [
        ("materials", "material_analysis", "1:1"),
        ("materials", "material_chunks", "1:N"),
        ("materials", "node_material_links", "1:N"),
        ("knowledge_topics", "knowledge_nodes", "1:N"),
        ("knowledge_nodes", "node_material_links", "1:N"),
        ("knowledge_nodes", "knowledge_edges", "1:N"),
        ("reconstruction_jobs", "knowledge_versions", "1:N"),
        ("knowledge_versions", "knowledge_topics", "1:N"),
    ]:
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        d.line((x1, y1, x2, y2), fill="#333333", width=2)
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        d.rectangle((mx - 25, my - 12, mx + 25, my + 12), fill="white")
        d.text((mx - 18, my - 11), lab, font=pf(14, "song"), fill="#111111")
    return save(img, "er.png")


def flow(name, title, nodes):
    img = Image.new("RGB", (1300, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 40), title, font=pf(28, "hei"), fill="#111111")
    y = 130
    for i, n in enumerate(nodes):
        d.rounded_rectangle((250, y, 1050, y + 60), radius=3, fill="#FFFFFF", outline="#222222", width=2)
        centered(d, (250, y, 1050, y + 60), f"{i+1}. {n}", pf(20, "song"))
        if i < len(nodes) - 1:
            arrow(d, 650, y + 60, 650, y + 105)
        y += 105
    return save(img, name)


def figures():
    IMG.mkdir(parents=True, exist_ok=True)
    return {
        "cover": cover_image(),
        "usecase": usecase(),
        "seq_material": seq("seq_material.png", "新增素材并解析时序图", ["用户", "前端", "API", "AI服务", "数据库"], [
            ("用户", "前端", "输入原始素材", False),
            ("前端", "API", "提交素材", False),
            ("API", "数据库", "保存 raw_content", False),
            ("前端", "API", "请求 AI 解析", False),
            ("API", "AI服务", "发送解析 Prompt", False),
            ("AI服务", "API", "返回结构化 JSON", True),
            ("API", "数据库", "写入解析结果", False),
            ("API", "前端", "返回解析状态", True),
        ]),
        "seq_system": seq("seq_system.png", "一键体系化重构时序图", ["用户", "工作台", "体系化API", "AI服务", "数据库"], [
            ("用户", "工作台", "点击一键体系化", False),
            ("工作台", "体系化API", "POST /api/systematize", False),
            ("体系化API", "数据库", "读取 analyzed 素材", False),
            ("体系化API", "AI服务", "生成知识体系 JSON", False),
            ("AI服务", "体系化API", "返回主题/节点", True),
            ("体系化API", "数据库", "写入 topics/nodes/edges", False),
            ("体系化API", "工作台", "返回版本号", True),
        ]),
        "architecture": architecture(),
        "er": er(),
        "flow": flow("overall_flow.png", "序知软件总体流程图", ["用户登录系统", "自由录入原始素材", "AI 完成单条解析", "用户触发一键体系化", "生成知识主题、节点与关联", "用户查看、问答或导出知识库"]),
        "test": flow("test_flow.png", "序知测试流程图", ["配置环境变量和数据库迁移", "执行 TypeScript 类型检查", "执行 Next.js 生产构建", "验证核心页面与接口路由", "记录测试结果和改进建议"]),
    }


def title_page(doc, title):
    para(doc, "", after=150)
    para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体", size=18, bold=True, after=0)
    doc.add_page_break()


def toc_page(doc, entries):
    para(doc, "目  录", align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体", size=16, bold=True, after=18)
    for item in entries:
        para(doc, item, font="宋体", size=10.5, after=1)
    doc.add_page_break()


def add_pic(doc, path, caption, width=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, font="宋体", size=10.5, after=6)


def requirements_acquisition(doc, fig):
    title_page(doc, "“序知”AI 个人体系化知识库\n软件需求获取文档")
    toc_page(doc, [
        "1.项目定位.......................................................... 1",
        "1.1 项目背景 ....................................................................................................................... 1",
        "1.2 项目目的 ....................................................................................................................... 1",
        "1.3 目标用户 ....................................................................................................................... 1",
        "2.成员分工.......................................................... 2",
        "3.软件需求.......................................................... 2",
        "3.1 功能需求 ....................................................................................................................... 2",
        "3.2 非功能需求 ................................................................................................................... 3",
        "4.软件需求描述分析.................................................. 3",
        "4.1 用户用例 ....................................................................................................................... 4",
    ])
    heading(doc, "1.项目定位", 1)
    heading(doc, "1.1 项目背景", 2)
    para(doc, "在个人学习和工作场景中，用户每天都会产生大量碎片化内容，包括课程笔记、会议记录、临时灵感、网页资料、读书摘录、经验复盘和问题清单。普通笔记软件通常要求用户自行分类、命名、排序和整理，长期使用后容易出现素材堆积、目录混乱、重复内容多、复用困难等问题。")
    para(doc, "“序知”的开发目标是解决“积累了知识但永远杂乱无体系”的痛点。系统允许用户以完全无序的方式录入原始素材，由 AI 完成深度解析、自动归类、去重、排序、层级化和知识网络构建，使用户逐步形成可长期使用的个人知识体系。")
    heading(doc, "1.2 项目目的", 2)
    bullet(doc, "实现原始素材的低门槛入库，用户不需要提前整理格式、分类和逻辑。")
    bullet(doc, "通过 AI 对单条素材进行核心含义、有效信息、主题和关键词提取。")
    bullet(doc, "通过一键体系化功能，将多条无序素材重构为“一级主题—二级分支—知识节点”的结构。")
    bullet(doc, "保留每个知识节点的来源素材引用，保证 AI 输出可追溯、可检查。")
    bullet(doc, "提供知识问答、知识关联、版本历史和数据导出能力，使知识库持续演化。")
    heading(doc, "1.3 目标用户", 2)
    para(doc, "核心用户包括学生、职场用户、创作者和终身学习者。学生主要用于课程知识整理和复习；职场用户主要用于会议纪要、项目复盘和工作经验沉淀；创作者主要用于灵感和素材聚合；终身学习者则用于跨领域资料的长期积累和关联。")
    heading(doc, "2.成员分工", 1)
    bullet(doc, "需求分析：根据“输入无序碎片，输出终身知识体系”的产品定位梳理功能边界。")
    bullet(doc, "系统设计：设计前后端一体化架构、Supabase 数据库结构和 AI 处理流程。")
    bullet(doc, "前端实现：完成登录、工作台、素材页、知识页、问答页和设置页。")
    bullet(doc, "后端实现：完成素材、解析、体系化、知识网络、问答和导出接口。")
    bullet(doc, "测试与文档：完成类型检查、构建验证、测试用例和用户说明文档。")
    heading(doc, "3.软件需求", 1)
    heading(doc, "3.1 功能需求", 2)
    table(doc, ["模块", "核心功能"], [
        ["用户认证", "- Magic Link 登录；\n- 登录后才能访问工作台、素材、知识和设置页面。"],
        ["素材管理", "- 新增、查看、编辑、删除原始素材；\n- 支持文本、OCR、音频转写、PDF/Word 解析等素材来源。"],
        ["AI 解析", "- 提取核心含义、有效信息、冗余信息、主题、知识类型和关键词；\n- 将长文本切分并生成向量。"],
        ["体系化重构", "- 一键生成主题、分支、知识节点；\n- 自动保留来源素材引用并生成版本快照。"],
        ["知识网络", "- 展示知识树、节点详情、来源素材和相关节点；\n- 通过 knowledge_edges 维护节点关系。"],
        ["知识问答", "- 基于用户知识库回答问题；\n- 支持将问答结果作为素材回存的扩展路径。"],
        ["数据与设置", "- 支持模型选择、API Key 配置、JSON/Markdown 导出和隐私说明。"],
    ], [3.2, 12.2])
    heading(doc, "3.2 非功能需求", 2)
    bullet(doc, "性能需求：核心页面应具备较快响应速度；体系化任务对单次处理素材数量进行控制。")
    bullet(doc, "安全需求：用户数据通过 Supabase Auth 和 RLS 隔离，服务端接口再次按 user_id 过滤。")
    bullet(doc, "可靠性需求：原始素材永远保留，AI 解析失败不得影响 raw_content。")
    bullet(doc, "可维护性需求：前后端使用 TypeScript，数据库迁移文件按功能拆分。")
    bullet(doc, "扩展性需求：素材层、解析层和知识网络层分离，便于后续增加增量重构和多模态素材。")
    heading(doc, "4.软件需求描述分析", 1)
    heading(doc, "4.1 用户用例", 2)
    add_pic(doc, fig["usecase"], "图 1 用户用例图", 15.7)
    for title, lines in [
        ("4.1.1 登录", ["主要参与者：用户。", "系统：验证用户身份并创建会话。", "前置条件：用户已拥有可接收 Magic Link 的邮箱。", "后置条件：用户进入主界面。", "基本交互序列：用户输入邮箱，系统发送登录链接，用户完成验证后进入工作台。"]),
        ("4.1.2 新增原始素材", ["主要参与者：用户。", "系统：保存原始素材并标记处理状态。", "前置条件：用户已登录。", "后置条件：materials 表新增记录。", "基本交互序列：用户输入标题和内容，点击保存，系统写入 raw_content。"]),
        ("4.1.3 AI 单条解析", ["主要参与者：用户、AI 服务。", "系统：读取素材并调用 AI 解析。", "前置条件：素材存在且属于当前用户。", "后置条件：生成 material_analysis 和 material_chunks。", "异常情况：AI 调用失败时素材状态标记为 failed。"]),
        ("4.1.4 一键体系化", ["主要参与者：用户、AI 服务、数据库。", "系统：将已解析素材重构为知识体系。", "前置条件：存在 analyzed 状态素材。", "后置条件：生成 topics、nodes、links、edges 和 version。", "异常情况：无可用素材时返回提示，不改变原始素材。"]),
        ("4.1.5 查看知识体系", ["主要参与者：用户。", "系统：展示 AI 生成的主题树和节点详情。", "前置条件：用户已有体系化结果。", "后置条件：用户可查看节点内容、来源素材和相关节点。"]),
        ("4.1.6 知识问答", ["主要参与者：用户、AI 服务。", "系统：基于个人知识库上下文回答问题。", "前置条件：用户已登录并拥有素材或知识节点。", "后置条件：返回基于知识库的回答。"]),
    ]:
        heading(doc, title, 3)
        for line in lines:
            bullet(doc, line)


def analysis_doc(doc, fig):
    title_page(doc, "“序知”AI 个人体系化知识库\n软件需求分析文档")
    toc_page(doc, [
        "1.运行环境 .......................................................... 1",
        "1.1 设备 ................................................................................................................................. 1",
        "1.2 接口 ................................................................................................................................. 1",
        "2.软件需求 .......................................................... 1",
        "2.1 功能需求 ........................................................................................................................ 1",
        "2.2 非功能需求 ................................................................................................................... 2",
        "2.3 软件总体流程 .............................................................................................................. 2",
        "2.4 软件体系结构 .............................................................................................................. 3",
        "3.需求分析——功能建模 ............................................. 4",
    ])
    heading(doc, "1.运行环境", 1)
    heading(doc, "1.1 设备", 2)
    para(doc, "浏览器要求：Chrome、Edge、Firefox、Safari 等现代浏览器。开发环境使用 Windows 本地目录 D:\\OrdKnow，生产环境可部署到 Vercel 或 Node.js 服务器。")
    heading(doc, "1.2 接口", 2)
    para(doc, "系统依赖 Supabase Auth、PostgreSQL、Storage、pgvector，以及 OpenAI 兼容 AI Provider、Embedding Provider、OCR 和音频转写接口。")
    heading(doc, "2.软件需求", 1)
    heading(doc, "2.1 功能需求", 2)
    table(doc, ["模块", "核心功能"], [
        ["用户认证", "- Magic Link 登录；\n- 受保护页面访问控制。"],
        ["素材管理", "- 素材新增、编辑、删除、搜索和状态筛选；\n- 支持多来源素材。"],
        ["AI 分析", "- 单条素材结构化解析；\n- 生成 chunk 和 embedding。"],
        ["体系重构", "- 生成知识主题、节点、引用和版本；\n- 建立最小知识网络。"],
        ["问答导出", "- 基于知识库问答；\n- 支持 JSON 和 Markdown 导出。"],
    ], [3.2, 12.2])
    heading(doc, "2.2 非功能需求", 2)
    bullet(doc, "安全性：RLS 策略确保用户只能读写自己的数据。")
    bullet(doc, "可维护性：核心业务类型集中定义在 src/types/index.ts。")
    bullet(doc, "兼容性：基于 Web 浏览器访问，适配 PC 端主要页面。")
    bullet(doc, "可靠性：构建验证通过，已执行 npm run lint 和 npm run build。")
    heading(doc, "2.3 软件总体流程", 2)
    add_pic(doc, fig["flow"], "图 2 软件总体流程图", 13.8)
    heading(doc, "2.4 软件体系结构", 2)
    add_pic(doc, fig["architecture"], "图 3 软件体系结构图", 15.7)
    heading(doc, "3.需求分析——功能建模", 1)
    heading(doc, "3.1 素材管理模块", 2)
    para(doc, "用户在工作台输入任意原始素材，系统保存原文并触发解析。原始素材层不被 AI 覆盖，是整个知识库的事实来源。")
    add_pic(doc, fig["seq_material"], "图 4 新增素材并解析时序图", 15.7)
    heading(doc, "3.2 体系化重构模块", 2)
    para(doc, "用户点击“一键体系化”后，系统读取已解析素材，调用 AI 生成体系化 JSON，并写入主题、节点、引用、关系和版本快照。")
    add_pic(doc, fig["seq_system"], "图 5 一键体系化重构时序图", 15.7)


def design_doc(doc, fig):
    title_page(doc, "“序知”AI 个人体系化知识库\n软件概要设计文档")
    toc_page(doc, [
        "1.系统总体设计 ...................................................... 1",
        "1.1 设计目标 .................................................................................................................... 1",
        "1.2 架构设计 .................................................................................................................... 1",
        "2.数据库设计 ........................................................ 2",
        "2.1 数据库结构 ................................................................................................................ 2",
        "2.2 数据表说明 ................................................................................................................ 3",
        "3.接口设计 .......................................................... 4",
        "4.模块设计 .......................................................... 5",
    ])
    heading(doc, "1.系统总体设计", 1)
    heading(doc, "1.1 设计目标", 2)
    para(doc, "系统设计遵循“原始素材永久保留、AI 只在理解层和知识网络层工作、所有输出可追溯到来源素材”的原则。")
    heading(doc, "1.2 架构设计", 2)
    add_pic(doc, fig["architecture"], "图 6 系统总体架构图", 15.7)
    heading(doc, "2.数据库设计", 1)
    heading(doc, "2.1 数据库结构", 2)
    add_pic(doc, fig["er"], "图 7 核心数据表关系图", 15.7)
    heading(doc, "2.2 数据表说明", 2)
    table(doc, ["数据表", "作用", "关键字段"], [
        ["materials", "保存用户原始素材。", "id、user_id、raw_content、source_type、status"],
        ["material_analysis", "保存 AI 单条解析结果。", "material_id、core_meaning、topics、keywords"],
        ["material_chunks", "保存素材分块和向量。", "content、chunk_index、embedding"],
        ["knowledge_topics", "保存一级主题和二级分支。", "parent_id、title、level、sort_order"],
        ["knowledge_nodes", "保存体系化知识节点。", "topic_id、title、content、node_type"],
        ["node_material_links", "保存节点到素材的引用。", "node_id、material_id、relevance_score"],
        ["knowledge_edges", "保存知识节点关系。", "source_node_id、target_node_id、edge_type"],
        ["knowledge_versions", "保存体系化版本快照。", "version_number、snapshot、summary"],
    ], [3.8, 4.8, 6.8])
    heading(doc, "3.接口设计", 1)
    table(doc, ["接口", "方法", "功能"], [
        ["/api/materials", "GET/POST", "查询和新增素材。"],
        ["/api/materials/[id]", "GET/PATCH/DELETE", "素材详情、编辑和删除。"],
        ["/api/analyze", "POST", "AI 单条解析。"],
        ["/api/systematize", "POST/GET", "触发体系化和查询任务。"],
        ["/api/knowledge", "GET", "获取知识主题树。"],
        ["/api/qa", "POST", "基于知识库问答。"],
        ["/api/export", "GET", "导出 JSON 数据。"],
        ["/api/export/markdown", "GET", "导出 Markdown 知识体系。"],
    ], [5.0, 2.5, 7.9])
    heading(doc, "4.模块设计", 1)
    bullet(doc, "素材模块：负责原始素材 CRUD、状态筛选和详情展示。")
    bullet(doc, "AI 模块：负责统一封装模型请求、结构化解析和体系化 Prompt。")
    bullet(doc, "知识模块：负责主题树、节点详情、来源引用和知识边查询。")
    bullet(doc, "问答模块：负责服务端检索用户知识上下文并调用 AI 生成回答。")
    bullet(doc, "设置模块：负责模型选择、API Key 配置、隐私说明和数据导出。")


def testing_manual(doc, fig):
    title_page(doc, "“序知”AI 个人体系化知识库\n测试与用户使用说明书")
    toc_page(doc, [
        "1.测试目标 .......................................................... 1",
        "2.测试环境 .......................................................... 1",
        "3.测试用例 .......................................................... 2",
        "4.运行环境 .......................................................... 3",
        "5.用户使用说明 ...................................................... 4",
    ])
    heading(doc, "1.测试目标", 1)
    para(doc, "测试目标是验证序知的核心闭环是否可用：用户能够登录系统，录入原始素材，AI 能完成单条解析，系统能进行一键体系化，用户能查看知识体系、进行知识问答并导出数据。")
    heading(doc, "2.测试环境", 1)
    table(doc, ["项目", "配置"], [
        ["操作系统", "Windows，本地开发目录 D:\\OrdKnow"],
        ["前端框架", "Next.js 16.2.7，React 19.2.7，TypeScript 6.0.3"],
        ["数据库", "Supabase PostgreSQL，启用 pgvector"],
        ["验证命令", "npm run lint；npm run build"],
        ["验证结果", "TypeScript 检查通过；Next.js 生产构建通过。"],
    ], [4.0, 11.4])
    add_pic(doc, fig["test"], "图 8 系统测试流程图", 13.5)
    heading(doc, "3.测试用例", 1)
    table(doc, ["编号", "测试内容", "操作步骤", "预期结果"], [
        ["1", "类型检查", "执行 npm run lint。", "tsc 无类型错误。"],
        ["2", "生产构建", "执行 npm run build。", "Next.js 编译成功并输出路由表。"],
        ["3", "素材新增", "登录后在工作台输入素材并保存。", "素材列表出现新记录。"],
        ["4", "AI 解析", "对 pending 素材触发解析。", "状态变为 analyzed 并显示解析结果。"],
        ["5", "一键体系化", "点击一键体系化按钮。", "生成主题、节点、引用和版本号。"],
        ["6", "知识问答", "在问答页输入问题。", "回答基于用户知识库上下文。"],
        ["7", "数据导出", "在设置页点击导出。", "下载 JSON 或 Markdown 文件。"],
    ], [1.2, 3.0, 5.8, 5.4])
    heading(doc, "4.运行环境", 1)
    heading(doc, "4.1 硬件运行环境", 2)
    para(doc, "开发端建议使用 8GB 以上内存的 PC；服务端可部署在支持 Node.js 的云平台，数据库使用 Supabase 托管。")
    heading(doc, "4.2 软件运行环境", 2)
    bullet(doc, "Node.js 与 npm：用于安装依赖、启动开发服务和执行构建。")
    bullet(doc, "Supabase：用于认证、数据库、存储和 pgvector。")
    bullet(doc, "AI Provider：用于素材解析、体系化重构和知识问答。")
    heading(doc, "5.用户使用说明", 1)
    for title, lines in [
        ("5.1 登录界面", ["用户输入邮箱。", "点击登录按钮。", "根据邮箱 Magic Link 完成验证。"]),
        ("5.2 工作台界面", ["左侧查看原始素材列表。", "中间输入或编辑素材。", "右侧查看 AI 体系化知识树。"]),
        ("5.3 素材界面", ["查看所有素材。", "按状态筛选 pending、analyzing、analyzed、failed。", "打开素材详情查看 AI 解析结果。"]),
        ("5.4 知识体系界面", ["查看一级主题和二级分支。", "点击节点查看内容。", "查看节点来源素材和相关节点。"]),
        ("5.5 问答界面", ["输入关于个人知识库的问题。", "系统基于知识库上下文回答。", "优质问答结果可作为后续素材回存。"]),
        ("5.6 设置界面", ["选择 AI 模型。", "配置用户自带 API Key。", "导出个人知识库数据。"]),
    ]:
        heading(doc, title, 2)
        for i, line in enumerate(lines, 1):
            numbered(doc, line, i)


def build():
    OUT.mkdir(exist_ok=True)
    if IMG.exists():
        shutil.rmtree(IMG)
    fig = figures()
    doc = Document()
    configure(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 控制封面图在正文区域内，避免 Word 因图片过高自动挤出一张空白页。
    p.add_run().add_picture(str(fig["cover"]), width=Cm(15.4))
    doc.add_page_break()
    requirements_acquisition(doc, fig)
    analysis_doc(doc, fig)
    design_doc(doc, fig)
    testing_manual(doc, fig)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
